"""
DOCX Generator - Converts markdown article content to a Word document
with embedded images and proper formatting.
"""

import os
import re
import httpx
from io import BytesIO
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _parse_markdown_to_blocks(content: str) -> list:
    """Parse markdown content into structured blocks."""
    blocks = []
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Title
        if line.startswith('TITLE:'):
            blocks.append({'type': 'title', 'text': line.replace('TITLE:', '').strip()})
            i += 1
            continue
        
        # H2
        if line.startswith('## '):
            blocks.append({'type': 'h2', 'text': line[3:].strip()})
            i += 1
            continue
        
        # H3
        if line.startswith('### '):
            blocks.append({'type': 'h3', 'text': line[4:].strip()})
            i += 1
            continue
        
        # Image
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if img_match:
            blocks.append({'type': 'image', 'alt': img_match.group(1), 'src': img_match.group(2)})
            i += 1
            continue
        
        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append({'type': 'list', 'items': items})
            continue
        
        # Numbered list
        num_match = re.match(r'^\d+\.\s', line)
        if num_match:
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                items.append(re.sub(r'^\d+\.\s', '', lines[i].strip()))
                i += 1
            blocks.append({'type': 'numbered_list', 'items': items})
            continue
        
        # Regular paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#') and not lines[i].strip().startswith('![') and not lines[i].strip().startswith('- ') and not lines[i].strip().startswith('* ') and not re.match(r'^\d+\.\s', lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        if para_lines:
            blocks.append({'type': 'paragraph', 'text': ' '.join(para_lines)})
            continue
        
        i += 1
    
    return blocks


def _add_rich_text(paragraph, text: str):
    """Add text with inline markdown formatting (bold, italic, links) to a paragraph."""
    # Pattern to match **bold**, *italic*, and [text](url)
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|\[([^\]]+)\]\(([^)]+)\))'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4) and match.group(5):  # Link
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # Add hyperlink
            run = paragraph.add_run(match.group(4))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True
            
            # Create actual hyperlink in the document
            r_id = paragraph.part.relate_to(
                match.group(5),
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            )
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            hyperlink.append(run._element)
            paragraph._element.append(hyperlink)
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _resolve_image_path(src: str, order_id: str) -> Optional[str]:
    """Resolve an image src to a local file path."""
    # Strip query string (cache busters like ?t=...)
    src = src.split('?')[0]
    
    # If it's an API URL like /api/v1/images/{order_id}/image_1.jpg
    if src.startswith('/api/v1/images/'):
        # Extract path: data/images/{order_id}/image_N.jpg
        parts = src.replace('/api/v1/images/', '').split('/')
        if len(parts) >= 2:
            local_path = f"data/images/{parts[0]}/{parts[1]}"
            if os.path.exists(local_path):
                return local_path
    
    # If it's already a local path
    if os.path.exists(src):
        return src
    
    return None


def generate_docx(
    article_content: str,
    order_id: str,
    output_path: Optional[str] = None,
) -> str:
    """
    Generate a Word document from markdown article content.
    
    Args:
        article_content: Markdown article content
        order_id: Order ID (for resolving image paths)
        output_path: Optional output file path
        
    Returns:
        Path to the generated .docx file
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Parse markdown
    blocks = _parse_markdown_to_blocks(article_content)
    
    for block in blocks:
        if block['type'] == 'title':
            p = doc.add_heading(block['text'], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        elif block['type'] == 'h2':
            doc.add_heading(block['text'], level=2)
            
        elif block['type'] == 'h3':
            doc.add_heading(block['text'], level=3)
            
        elif block['type'] == 'paragraph':
            p = doc.add_paragraph()
            _add_rich_text(p, block['text'])
            
        elif block['type'] == 'list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Bullet')
                _add_rich_text(p, item)
                
        elif block['type'] == 'numbered_list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Number')
                _add_rich_text(p, item)
                
        elif block['type'] == 'image':
            img_path = _resolve_image_path(block['src'], order_id)
            if img_path:
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    # Add caption
                    if block['alt']:
                        caption = doc.add_paragraph(block['alt'])
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.runs[0].font.size = Pt(9) if caption.runs else None
                        caption.runs[0].font.italic = True if caption.runs else None
                        caption.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66) if caption.runs else None
                except Exception as e:
                    # If image fails, add placeholder text
                    p = doc.add_paragraph(f"[Image: {block['alt']}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(f"[Image: {block['alt']}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    if not output_path:
        output_dir = Path(f"data/docx/{order_id}")
        output_dir.mkdir(parents=True, exist_ok=True)
        # Use title for filename
        title = "article"
        for block in blocks:
            if block['type'] == 'title':
                # Clean title for filename
                title = re.sub(r'[^\w\s-]', '', block['text'])[:60].strip().replace(' ', '_')
                break
        output_path = str(output_dir / f"{title}.docx")
    
    doc.save(output_path)
    return output_path
